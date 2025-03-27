import tkinter as tk
from tkinter import filedialog, messagebox
from tkinter import ttk
import os
from docx import Document
import re
import openpyxl
import threading

def substituir_texto(arquivo, texto_antigo, texto_novo):
    doc = Document(arquivo)
    
    def substituir_em_runs(runs):
        texto_completo = "".join(run.text for run in runs)
        novo_texto = re.sub(re.escape(texto_antigo), texto_novo, texto_completo, flags=re.IGNORECASE)
        if texto_completo != novo_texto:
            for i, run in enumerate(runs):
                if i == 0:
                    run.text = novo_texto
                else:
                    run.text = ""
    
    for paragrafo in doc.paragraphs:
        substituir_em_runs(paragrafo.runs)
    
    for tabela in doc.tables:
        for linha in tabela.rows:
            for celula in linha.cells:
                for paragrafo in celula.paragraphs:
                    substituir_em_runs(paragrafo.runs)
    
    doc.save(arquivo)

def processar_texto(arquivo_xlsx, caminho_arquivos, progress_var, status_var):
    workbook = openpyxl.load_workbook(arquivo_xlsx)
    sheet = workbook.active
    substituicoes = []
    
    for row in sheet.iter_rows(min_row=1, values_only=True):
        if len(row) >= 2:
            texto_antigo, texto_novo = row[0], row[1]
            substituicoes.append((texto_antigo, texto_novo))
    
    arquivos_processados = 0
    total_arquivos = len([f for f in os.listdir(caminho_arquivos) if f.endswith(".docx")])
    
    for arquivo in os.listdir(caminho_arquivos):
        if arquivo.endswith(".docx"):
            caminho_completo = os.path.join(caminho_arquivos, arquivo)
            status_var.set(f"Processando: {arquivo}")
            for texto_antigo, texto_novo in substituicoes:
                substituir_texto(caminho_completo, texto_antigo, texto_novo)
            arquivos_processados += 1
            progress_var.set(int((arquivos_processados / total_arquivos) * 100))
    
    workbook.close()
    status_var.set("Processamento concluído!")
    return arquivos_processados

class Application(tk.Tk):
    def __init__(self):
        super().__init__()
        self.title("Substituição de Texto em Documentos Word")
        self.geometry("500x300")
        
        self.xlsx_path = tk.StringVar()
        self.docx_folder = tk.StringVar()
        self.progress_var = tk.IntVar()
        self.status_var = tk.StringVar()
        
        self.create_widgets()
    
    def create_widgets(self):
        # Arquivo Excel
        tk.Label(self, text="Arquivo Excel (.xlsx):").grid(row=0, column=0, padx=5, pady=5, sticky="w")
        tk.Entry(self, textvariable=self.xlsx_path, width=50).grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        tk.Button(self, text="Procurar", command=self.browse_xlsx).grid(row=0, column=2, padx=5, pady=5, sticky="w")
        
        # Pasta com documentos Word
        tk.Label(self, text="Pasta com documentos Word:").grid(row=1, column=0, padx=5, pady=5, sticky="w")
        tk.Entry(self, textvariable=self.docx_folder, width=50).grid(row=1, column=1, padx=5, pady=5, sticky="ew")
        tk.Button(self, text="Procurar", command=self.browse_folder).grid(row=1, column=2, padx=5, pady=5, sticky="w")
        
        # Botão Iniciar Processamento
        tk.Button(self, text="Iniciar Processamento", command=self.start_processing).grid(row=2, column=0, columnspan=3, pady=10)
        
        # Barra de Progresso
        self.progress_bar = ttk.Progressbar(self, variable=self.progress_var, maximum=100)
        self.progress_bar.grid(row=3, column=0, columnspan=3, padx=5, pady=5, sticky="ew")
        
        # Label de Status
        tk.Label(self, textvariable=self.status_var).grid(row=4, column=0, columnspan=3, pady=5)

        # Configurar pesos das colunas para expandir o Entry
        self.grid_columnconfigure(1, weight=1)
    
    def browse_xlsx(self):
        filename = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])
        if filename:
            self.xlsx_path.set(filename)
    
    def browse_folder(self):
        folder = filedialog.askdirectory()
        if folder:
            self.docx_folder.set(folder)
    
    def start_processing(self):
        xlsx_file = self.xlsx_path.get()
        docx_folder = self.docx_folder.get()
        
        if not xlsx_file or not docx_folder:
            messagebox.showerror("Erro", "Por favor, selecione o arquivo Excel e a pasta com os documentos Word.")
            return
        
        self.progress_var.set(0)
        self.status_var.set("Iniciando processamento...")
        
        threading.Thread(target=self.process_files, args=(xlsx_file, docx_folder), daemon=True).start()
    
    def process_files(self, xlsx_file, docx_folder):
        try:
            arquivos_processados = processar_texto(xlsx_file, docx_folder, self.progress_var, self.status_var)
            messagebox.showinfo("Concluído", f"Processamento concluído!\nTotal de arquivos processados: {arquivos_processados}")
        except Exception as e:
            messagebox.showerror("Erro", f"Ocorreu um erro durante o processamento: {str(e)}")

if __name__ == "__main__":
    app = Application()
    app.mainloop()
